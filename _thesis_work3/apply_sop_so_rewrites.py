"""
Apply accessible-register SOP/SO rewrites to FINAL-THESIS-3-PAPER-1 (2).docx.

The current paragraphs (verified by scan):
  SOPs : indices 174..178  (List Paragraph style, under "Statement of the Problem")
  SOs  : indices 182..186  (List Paragraph style, under "Objectives of the Study")

Rewrites are sourced verbatim from _thesis_work3/isolation-audit.md blocks A and B.

Strategy:
  For each target paragraph, clear all runs and add ONE new run with the new
  text. This preserves the paragraph's STYLE (List Paragraph) and numbering,
  but does not preserve mid-sentence inline formatting -- inspection of the
  original paragraphs showed plain body text, so no formatting is lost.

Safety:
  A backup copy was made before this script runs:
    FINAL-THESIS-3-PAPER-1 (2).backup-before-sop-rewrite.docx
"""

from pathlib import Path
import docx

DOC_PATH = Path("FINAL-THESIS-3-PAPER-1 (2).docx")

# Index -> new text. Maps to the accessible-register rewrites from
# _thesis_work3/isolation-audit.md blocks A and B.
NEW_TEXT = {
    # --- Statement of the Problem ---
    174: (
        "How can the system teach design patterns in a way that lets the user "
        "try each pattern on their own C++ code immediately after the lesson, "
        "instead of only reading about it?"
    ),
    175: (
        "How can the system examine a user's C++ code to find which design "
        "patterns are present, without changing the user's original code, and "
        "allow the team to add support for new patterns later without "
        "rebuilding the system?"
    ),
    176: (
        "How can the system present its findings in a way that shows the user "
        "which lines of their code triggered each detected pattern and why, "
        "so the user can verify the finding for themselves?"
    ),
    177: (
        "How can the system write per-class documentation that combines "
        "findings the system verified from the user's code with explanations "
        "written by AI, while clearly labelling which sentences come from "
        "which source so the user can trust the documentation?"
    ),
    178: (
        "How well does the system help DEVCON Luzon participants identify, "
        "document, and explain design patterns in C++ code that someone else "
        "wrote — measured by comparing what they can do before and after "
        "using the system?"
    ),
    # --- Specific Objectives ---
    182: (
        "Create learning modules where each design-pattern lesson is paired "
        "with a hands-on check the system runs on the user's own C++ code, "
        "so the user sees the lesson applied to their work right after "
        "learning it."
    ),
    183: (
        "Build a C++ code-analysis feature that reads the user's source "
        "without modifying it, identifies which design patterns appear at the "
        "class level, and reads its pattern definitions from configuration "
        "files so new patterns can be added by editing files instead of "
        "rebuilding the system."
    ),
    184: (
        "Build a results view that lists the most likely pattern matches "
        "first, links each match to the exact lines of code that caused it, "
        "and explains in plain language what the system saw in those lines."
    ),
    185: (
        "Build a documentation feature that produces a written explanation "
        "for each analysed class, mixing facts the system verified against "
        "the user's code with sentences written by AI, with each AI-written "
        "sentence visibly marked so the user can tell the two apart."
    ),
    186: (
        "Evaluate the system with DEVCON Luzon participants by measuring, "
        "before and after they use the system, how accurately they can "
        "identify design patterns in unfamiliar C++ code, how completely "
        "they can document those patterns, and how clearly they can explain "
        "why each pattern is there."
    ),
}


def replace_paragraph_text(paragraph, new_text):
    """Replace all text in `paragraph` with `new_text`, preserving the
    paragraph's style (List Paragraph numbering) and the first run's font
    properties.
    """
    # Capture the first run's font as a template so the replacement inherits
    # font family / size / colour from the original first run.
    runs = paragraph.runs
    template_run = runs[0] if runs else None

    # Delete every run by removing its underlying <w:r> element.
    for run in list(runs):
        run._element.getparent().remove(run._element)

    # Add a fresh run with the new text.
    new_run = paragraph.add_run(new_text)

    # Inherit font properties from the original template run.
    if template_run is not None:
        new_run.font.name = template_run.font.name
        new_run.font.size = template_run.font.size
        new_run.font.bold = template_run.font.bold
        new_run.font.italic = template_run.font.italic
        new_run.font.underline = template_run.font.underline
        if template_run.font.color and template_run.font.color.rgb:
            new_run.font.color.rgb = template_run.font.color.rgb


def main():
    if not DOC_PATH.exists():
        raise SystemExit(f"docx not found: {DOC_PATH.resolve()}")

    doc = docx.Document(str(DOC_PATH))

    # Sanity check: confirm the target paragraphs still match the expected
    # style + opening words. If something has shifted, refuse to write and
    # report instead -- the backup file is the recovery path.
    EXPECTED_OPENINGS = {
        174: "How can the proposed system provide learning support",
        175: "How can the proposed system analyze C++ source code",
        176: "How can the proposed system present detected design-pattern",
        177: "How can documentation-oriented outputs help users",
        178: "How useful is the proposed system in supporting",
        182: "Develop learning modules that support users",
        183: "Develop a C++ source-code analysis feature",
        184: "Present detected design-pattern evidence",
        185: "Generate documentation-oriented outputs",
        186: "Evaluate the usefulness of the proposed system",
    }

    problems = []
    for idx, expected_prefix in EXPECTED_OPENINGS.items():
        if idx >= len(doc.paragraphs):
            problems.append(f"  [{idx}] paragraph index out of range")
            continue
        actual = doc.paragraphs[idx].text.strip()
        if not actual.startswith(expected_prefix):
            problems.append(
                f"  [{idx}] expected start '{expected_prefix[:60]}...' "
                f"but found '{actual[:60]}...'"
            )

    if problems:
        print("ABORT: target paragraphs no longer match the expected layout.")
        print("The docx structure has shifted since the scan. No edits made.")
        print("Issues:")
        for p in problems:
            print(p)
        raise SystemExit(1)

    # Apply rewrites.
    for idx, new_text in NEW_TEXT.items():
        replace_paragraph_text(doc.paragraphs[idx], new_text)
        print(f"  [{idx}] rewrote -> {new_text[:80]}...")

    doc.save(str(DOC_PATH))
    print(f"\nSaved: {DOC_PATH}")
    print(f"Backup: FINAL-THESIS-3-PAPER-1 (2).backup-before-sop-rewrite.docx")


if __name__ == "__main__":
    main()
